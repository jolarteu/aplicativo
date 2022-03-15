from django.shortcuts import render, redirect
from django.http import HttpResponse
from  datetime import datetime
from django.contrib.auth.decorators import login_required #pide iniciar seccion
from django.views.decorators.cache import cache_control, never_cache
from django.contrib.auth.mixins import  LoginRequiredMixin
from django.forms import modelformset_factory, inlineformset_factory
from django import forms
# from post.models import Post
# from post.forms import PostForm
from django.views.generic import ListView, DetailView, CreateView, UpdateView
from django.urls import reverse
from Homologaciones.models import Homologacion, pais, fabricante, tipo, referencia, estado, resultado, atributo_elemento_h, fabricante_pais
from Homologaciones.forms import paisForm
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from Dispositivos.models import dispositivo as dispositivo_id
from Dispositivos.models import dispositivo_atributo, atributo
from django.views.generic.edit import FormMixin
from Homologaciones.forms import  HomologacionForm, atributo_elemento_hForm, ReferenciaForms, paisForm, HomologacionUpdateForm
# Create your views here.
from django.http import HttpResponseRedirect
from Users.models import Profile
from django.contrib.auth.models import User
import io
from django.http import FileResponse
from PyPDF2 import PdfFileWriter, PdfFileReader
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx
from datetime import date

@login_required()
def home(request):
    return render(request, 'Homologaciones/home.html')

@login_required()
def homologaciones(request):

    return render(request, 'Homologaciones/homologaciones.html')


class  Homologacion_terminar(DetailView ,UpdateView):
    model=Homologacion
    template_name= 'Homologaciones/terminar.html'
    form_class=HomologacionUpdateForm

    queryset = Homologacion.objects.all()

    # def get_queryset(self):
    #         #self.user = User.objects.filter(profile=self.request.user)
    #         self.user= self.request.user
    #         return self.user

    def get_object(self, queryset=None):
        self.obj = super().get_object()
        return self.obj

    def obtener_dispositivo(self, *args, **kwargs):
        self.obj=self.get_object()
        self.refer=referencia.objects.get(pk=self.obj.refer.pk)
        self.dispositivo=dispositivo_id.objects.get(pk=self.refer.id_dispositivo.pk)
        return self.dispositivo

    def form_set(self, *args, **kwargs):
        self.paisFormSet = modelformset_factory(atributo_elemento_h, fields=('atributo','valor','document'),
        extra=self.cantidad(), widgets={
            'valor': forms.TextInput(attrs={'class': 'myfieldclass'}),
            })


        return self.paisFormSet


    def form_set_2(self, *args, **kwargs):
        self.paisFormSet = modelformset_factory(Homologacion, fields=('__all__'),
        )
        return self.paisFormSet

    def cantidad(self, **kwargs):
        extra=dispositivo_atributo.objects.filter(id_dispositivo=self.obtener_dispositivo()).count()
        return extra

    def lista(self, **kwargs):
        lista=dispositivo_atributo.objects.filter(id_dispositivo=self.obtener_dispositivo())
        lista=[(lista[i].pk, lista[i]) for i in range(len(lista))]
        return lista

    def autollenado(self, i, **kwargs):

        lista=dispositivo_atributo.objects.filter(id_dispositivo=self.obtener_dispositivo())
        self.obj = super().get_object()
        cantidad=atributo_elemento_h.objects.filter(atributo=lista[i], Homologacion=self.obj).count()
        if cantidad==0:
            try:
                busqueda=Homologacion.objects.filter(refer=self.obj.refer).order_by('-id')[1]
                auto=atributo_elemento_h.objects.filter(atributo=lista[i], Homologacion=busqueda).last().valor

            except:auto=""
        else:
            try:
                auto=atributo_elemento_h.objects.filter(atributo=lista[i], Homologacion=self.obj).last().valor
            except:
                auto=""
        # else:
        #     try:auto=Homologacion.objects.filter(refer=self.obj.refer)
        #     except:
        #         auto="aaa"

        return(str(auto))                  ####revisar

    def contexto(self, **kwargs):
        pass
        return self.formset

    def get_context_data(self, **kwargs):
        self.obtener_dispositivo()
        self.HomologacionFormSet = self.form_set_2()
        self.paisFormSet = self.form_set()
        formset2=self.HomologacionFormSet(queryset=Homologacion.objects.none())
        formset=self.paisFormSet(queryset=dispositivo_atributo.objects.none())
        for i in range(self.cantidad()):
            formset[i].fields['atributo'].choices = self.lista()
            formset[i].fields['atributo'].initial  = (self.lista()[i])[0]
            formset[i].fields['valor'].initial=self.autollenado(i)
            #formset[i].fields['valor'].initial=(self.lista()[i])[1]

        #formset[0].fields['atributo'].choices = [(1,4),(2,2),(3,3)]
    #    formset[0].fields['atributo'].choices = self.lista()
        context = super(Homologacion_terminar, self).get_context_data(**kwargs)
    #    context['formset'] = self.paisFormSet(queryset=dispositivo_atributo.objects.none())
        #context['formset2'] = formset2
        context['formset'] = formset
        context['title'] = 'terminar homologacion'
        #context['pk']=Homologacion.objects.get(pk=self.obj.pk)
        return context

    def post(self, request , *args, **kwargs):
        self.obj = super().get_object()
        self.HomologacionFormSet = self.form_set_2()
        self.paisFormSet = self.form_set()
        formset2 = self.HomologacionFormSet(request.POST)
    #    form1=self.form(request.POST)
        formset = self.paisFormSet(request.POST, request.FILES)
        form=HomologacionUpdateForm(request.POST, request.FILES, instance=self.obj)
    #    form = self.get_form()
        user = Profile.objects.get(pk=request.user.profile.pk)
        if formset.is_valid() and form.is_valid():
            return self.form_valid(formset,form, user)


    def form_valid(self, formset, form, user):

        self.obj = super().get_object()
        instances = formset.save(commit=False)
        f=form.save(commit=False)
        if (f.resultado==resultado.objects.get(pk='Sin terminar')):
            f.estado=estado.objects.get(pk='En curso')
        else:
            f.estado=estado.objects.get(pk='Cerrada')
        f.save()
        # f.profile=user
        # f.refer=referencia.objects.get(pk=sel.obj.refer)
        #instance2= form.save(commit=False)
        for instance in instances:
            instance.Homologacion=Homologacion.objects.get(pk=self.obj.pk)
            #instance.profile=Profile.objects.get(pk=2)          #colocar para detectar usuario IMPORTANTE
            instance.profile=user           #SE PUDO :)
            instance.save()

        #return super(Createpais, self).form_valid(form)
        #return HttpResponseRedirect('Homologaciones:home')
        return redirect(self.get_success_url())

    def get_success_url(self):
         return reverse('Homologaciones:consultar')             #llenar datos para la homologacion

class CreateHomologacion(FormMixin, DetailView):
    template_name = 'Homologaciones/crear.html'
    form_class = HomologacionForm
    form_class2= atributo_elemento_hForm
    model = Homologacion


    queryset = referencia.objects.all()
    slug_field = 'pk'
    slug_url_kwarg = 'pk'
    queryset = referencia.objects.all()
    context_object_name = 'pk'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['estados'] = queryset=estado.objects.all()
        context['resultados'] = queryset=resultado.objects.all()

        return context

    def homologacion(self, request, *args, **kwargs):
        self.object = self.get_object()
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)

    def form_valid(self, form):
        form.save()
        return super(CreateHomologacion, self).form_valid(form)


    def get_success_url(self):
        return reverse('Homologacion:home')

class HomologacionDetailView(LoginRequiredMixin, DetailView):

    template_name = 'Homologaciones/detail.html'

    queryset = referencia.objects.all()
    slug_field = 'pk'
    slug_url_kwarg = 'pk'
    queryset = referencia.objects.all()
    context_object_name = 'pk'

    def get_object(self, queryset=None):
        self.obj = super().get_object()
        return self.obj

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['atributos'] = queryset=dispositivo_atributo.objects.filter(id_dispositivo=self.obj.id_dispositivo)
        return context

class HomologacionListView(LoginRequiredMixin, DetailView):
    model = Homologacion
    template_name = 'Homologaciones/lista_homologaciones.html'

    queryset = referencia.objects.all()

    def get_object(self, queryset=None):
        self.obj = super().get_object()
        return self.obj

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['comparar']=estado.objects.get(pk='En curso')
        context['title'] = 'Lista homologaciones'
        context['object_list']=Homologacion.objects.filter(refer=self.obj.pk).order_by('id')
        return context          #ver homologaciones de cadad terminal

    def version_sw(self, pk, ref):
        atri_busqueda=atributo.objects.filter(name="SW").order_by('-id')[0]
        atri_busqueda_2=dispositivo_atributo.objects.get(id_dispositivo=ref.id_dispositivo,id_atributo_id=atri_busqueda)
        atri=atributo_elemento_h.objects.filter(Homologacion=pk, atributo=atri_busqueda_2).order_by('-id')[0]
        #dis=dispositivo_id.objects.get(pk=ref.id_dispositivo)

        return str(atri.valor)

    def version_os(self, pk, ref):
        atri_busqueda=atributo.objects.filter(name="Sistema operativo").order_by('-id')[0]
        atri_busqueda_2=dispositivo_atributo.objects.get(id_dispositivo=ref.id_dispositivo,id_atributo_id=atri_busqueda)
        atri=atributo_elemento_h.objects.filter(Homologacion=pk, atributo=atri_busqueda_2).order_by('-id')[0]
        #dis=dispositivo_id.objects.get(pk=ref.id_dispositivo)

        return str(atri.valor)

    def version_per(self, pk, ref):
        atri_busqueda=atributo.objects.filter(name="Personalizacion").order_by('-id')[0]
        atri_busqueda_2=dispositivo_atributo.objects.get(id_dispositivo=ref.id_dispositivo,id_atributo_id=atri_busqueda)
        atri=atributo_elemento_h.objects.filter(Homologacion=pk, atributo=atri_busqueda_2).order_by('-id')[0]
        #dis=dispositivo_id.objects.get(pk=ref.id_dispositivo)

        return str(atri.valor)

    def dual_sim(self, pk, ref):
        atri_busqueda=atributo.objects.filter(name="Dual sim").order_by('-id')[0]
        atri_busqueda_2=dispositivo_atributo.objects.get(id_dispositivo=ref.id_dispositivo,id_atributo_id=atri_busqueda)
        atri=atributo_elemento_h.objects.filter(Homologacion=pk, atributo=atri_busqueda_2).order_by('-id')[0]
        #dis=dispositivo_id.objects.get(pk=ref.id_dispositivo)
        if str(atri.valor)=="SI" or str(atri.valor)=="si" or str(atri.valor)=="x" or str(atri.valor)=="X":
            return "Es dual Sim"
        else:
            return "Es singel Sim"

    def perso(self, pk, ref):
        atri_busqueda=atributo.objects.filter(name="Personalizacion").order_by('-id')[0]
        atri_busqueda_2=dispositivo_atributo.objects.get(id_dispositivo=ref.id_dispositivo,id_atributo_id=atri_busqueda)
        atri=atributo_elemento_h.objects.filter(Homologacion=pk, atributo=atri_busqueda_2).order_by('-id')[0]
        #dis=dispositivo_id.objects.get(pk=ref.id_dispositivo)
        if str(atri.valor)=="open" or str(atri.valor)=="OPEN" or str(atri.valor)=="Openla" or str(atri.valor)=="No" or str(atri.valor)=="NO" or str(atri.valor)=="No" or str(atri.valor)=="no":
            return "No cuenta con Personalizacion de Movistar"
        else:
            return "descarga la customización de Movistar "

    def post(self, request, pk):
        today = date.today()
        pk=request.POST['key']
        object=self.get_object()
        ref=referencia.objects.filter(refer=object.refer).order_by('-pk')[0]

        pais_busqueda=pais.objects.filter(pk=ref.pais).order_by('-pk')[0]
        fabri=fabricante.objects.filter(pk=ref.fabricante).order_by('-pk')[0]
        Nombre=fabricante_pais.objects.filter(pais=pais_busqueda, fabricante=fabri).order_by('-id')[0]
        doc = docx.Document('original_vacio.docx')
        doc.add_paragraph(str(pais_busqueda.capital)+", "+str(today)+"\n\n"+"Señor(a)")
        doc.add_paragraph(str(Nombre.representante))
        doc.add_paragraph(str(fabri)+" "+str(pais_busqueda))
        doc.add_paragraph(str(pais_busqueda.capital)+"\n\n")
        doc.add_paragraph("Asunto: Aprobación Técnica de el/la "+str(ref.id_dispositivo)+" "+ str(ref.refer )+" ("+str(ref.name)+")\n\n")
        doc.add_paragraph("Luego de realizadas las pruebas de Ingeniería de RF, Conmutación y Usabilidad de la referencia: "+ str(ref.refer )+" ("+str(ref.name)+")\n\n")

        doc.add_paragraph('Versión Firmware/Software: '+self.version_sw(pk, ref),
                  style='Estilo4')
        doc.add_paragraph('Versión Personalización: '+self.version_per(pk, ref),
                  style='Estilo4')
        doc.add_paragraph('Sistema Operativo: '+self.version_os(pk, ref),
                  style='Estilo4')

        doc.add_paragraph('\nSe da concepto de APROBACIÓN de este terminal en la red 2G/3G/4G de Telefónica Móviles Colombia. se hacen las siguientes aclaraciones sobre los resultados del proceso de homologación de este modelo:',
                  )

        doc.add_paragraph('Son terminales para navegación WEB',
                style='Estilo4')
        doc.add_paragraph('Los terminales tienen funcionamiento ALWAYS ON.',
                style='Estilo4')
        doc.add_paragraph('Son terminales para navegación WEB',
                style='Estilo4')
        doc.add_paragraph(self.dual_sim(pk, ref),
                style='Estilo4')
        doc.add_paragraph(self.perso(pk, ref),
                style='Estilo4')

        doc.add_paragraph('\nBugs en revisión, se espera sean resueltos en la próxima versión de mantenimiento')

        x=" debe entregar los certificados que sean requeridos de manera global, regional y local que estén pendientes por recibir. Así mismo debe garantizar que la información plasmada en las fichas técnicas enviadas sea real y no tenga ningún error."
        x2="""Como es habitual, continuaremos con el proceso de pruebas sobre los lotes de producción para verificar la eliminación SIM LOCK y el cumplimiento de las condiciones de configuración evaluadas. La omisión de estos requerimientos será motivo de devolución de los terminales. Por lo tanto y teniendo en cuenta los puntos que están descritos, se aprueba la comercialización de esta referencia de terminal con la versión de software menciona, sobre la red de Movistar Colombia.
                \nCordialmente, Oscar José Carvajal
                \nJefe de Homologación de Terminales Móviles"""

        doc.add_paragraph('\n'+str(fabri)+x)
        doc.add_paragraph('\n'+str(fabri)+x2)

        # doc.add_paragraph('Hello world!')
        doc.save('salida.docx')

        # # buffer = io.BytesIO()
        #
        #
        # # See the ReportLab documentation for the full list of functionality.
        # p.drawString(100, 100, "Hello world.")
        #
        # # Close the PDF object cleanly, and we're done.
        # p.showPage()
        # p.save()
        #
        # # FileResponse sets the Content-Disposition header so that browsers
        # # present the option to save the file.
        # buffer.seek(0)
        # new_pdf = PdfFileReader(buffer)
        # existing_pdf = PdfFileReader(open("original.pdf", "rb"))
        # output = PdfFileWriter()
        #
        # page = existing_pdf.getPage(0)
        # page.mergePage(new_pdf.getPage(0))
        # output.addPage(page)
        # #
        # outputStream = open("destination.pdf", "wb")
        # output.write(outputStream)
        # outputStream.close()

        return FileResponse(open('salida.docx', 'rb'), as_attachment=True, filename='hello.docx')
        #return FileResponse(response, as_attachment=True, filename='hello.pdf')
#        return redirect(self.get_success_url_1(self.pk))



class detalles_homologacionDetailview(LoginRequiredMixin, DetailView):
    model = atributo_elemento_h
    template_name = 'Homologaciones/detalles_homologacion.html'

    queryset= Homologacion.objects.all()

    def get_object(self, queryset=None):
        self.obj = super().get_object()
        return self.obj

    def obtener_dispositivo(self, *args, **kwargs):
        self.obj=self.get_object()
        self.refer=referencia.objects.get(pk=self.obj.refer.pk)
        self.dispositivo=dispositivo_id.objects.get(pk=self.refer.id_dispositivo.pk)
        return self.dispositivo

    def cantidad(self, **kwargs):
        extra=dispositivo_atributo.objects.filter(id_dispositivo=self.obtener_dispositivo()).count()
        return extra

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = Homologacion.objects.filter(pk=self.obj.pk)
        cantidad=atributo_elemento_h.objects.filter(Homologacion=self.obj.pk).distinct('atributo').count()
        #context['object_list']=atributo_elemento_h.objects.filter(Homologacion=self.obj.pk).order_by('-id').distinct('atributo')
        context['object_list']=atributo_elemento_h.objects.filter(Homologacion=self.obj.pk).order_by('-id')[:cantidad]

        return context          #ver homologaciones de cadad terminal           #lista de atributos

class CategoryListView(LoginRequiredMixin, ListView):
    model = referencia
    template_name = 'Homologaciones/list2.html'



    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Lista homologaciones'

        return context          #lista de terminales


    def get_success_url_1(self, pk):
        # return reverse('Homologaciones:consultar_terminal/'+pk)
        return reverse('Homologaciones:consultar_terminal',kwargs={'pk': pk})           #asi se redirige bien :)

    def get_success_url_3(self, pk):
        # return reverse('Homologaciones:consultar_terminal/'+pk)
        return reverse('Homologaciones:terminar',kwargs={'pk': pk})           #asi se redirige bien :)


    def post(self, request):
        self.pk=request.POST['key']
        self.function=request.POST['function']

        if self.function=="plus":
            Homologacion.objects.create(
                refer=referencia.objects.get(pk=self.pk),
                profile=Profile.objects.get(pk=request.user.profile.pk),
                tipo=tipo.objects.get(tipo="Oficial")

            )
            return redirect(self.get_success_url_1(self.pk))

        elif self.function=="delete":
            referencia.objects.filter(pk=self.pk).delete()
            return redirect(self.get_success_url_2())

        elif self.function=="see":
            Homologacion_t=Homologacion.objects.filter(refer=self.pk).last().pk
            return redirect(self.get_success_url_3(Homologacion_t))


    def get_success_url_2(self):
         return reverse('Homologaciones:consultar')


    # def form_valid(self, form):
    #     self.pk = self.request.pk
    #     messages.success(self.request, 'Form submission successful')
    #     Homologacion.objects.create(
    #               refer=referencia.objects.get(pk=32),
    #               profile=Profile.objects.get(pk=request.user.profile.pk),
    #               tipo=tipo.objects.get(tipo="Oficial")
    #     )
    #     return redirect(self.get_success_url())

class Createreferencia(LoginRequiredMixin,SuccessMessageMixin, CreateView):

    # model = referencia
    # fields = [ 'profile','id_dispositivo', 'refer', 'pais', 'name',  'fabricante']
    form_class= ReferenciaForms
#    form_class2= atributo_elemento_hForm
    template_name = 'Homologaciones/new4.html'


    def form_valid(self, form):
        form.instance.user = self.request.user
        form.instance.profile = self.request.user.profile
        print(form)
        form.save()
        self.new_refer=form.save()
        messages.success(self.request, 'Form submission successful')
        Homologacion.objects.create(
                refer=referencia.objects.get(pk=self.new_refer.pk),
                profile=self.new_refer.profile,
                tipo=tipo.objects.get(tipo="Oficial")
        )
        return super(Createreferencia, self).form_valid(form)

    def get_success_url(self):
        return reverse('Homologaciones:consultar_terminal', kwargs={ "pk": self.new_refer.pk})
        #return render(request, 'Homologaciones/detail', 1)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['fabricantes'] = queryset=fabricante.objects.all()
        context['dispositivos'] = queryset=dispositivo_id.objects.all()
        return context

class Createpais(LoginRequiredMixin, SuccessMessageMixin, CreateView):

    model = pais
    form_class= paisForm
    template_name = 'Homologaciones/paises.html'



    def get_context_data(self, **kwargs):
        self.paisFormSet = modelformset_factory(atributo_elemento_h, fields='__all__',extra=2)
        context = super(Createpais, self).get_context_data(**kwargs)
        context['formset'] = self.paisFormSet(queryset=pais.objects.none())
        return context

    def post(self, request, *args, **kwargs):
        self.paisFormSet = modelformset_factory(atributo_elemento_h, fields='__all__',extra=2)
        formset = self.paisFormSet(request.POST)
        if formset.is_valid():
            return self.form_valid(formset)

    def form_valid(self, formset):
        instances = formset.save()
        for instance in instances:
            print("prubeaaaaaa")
            instance.save()
        #return super(Createpais, self).form_valid(form)
        #return HttpResponseRedirect('Homologaciones:home')
        return redirect(self.get_success_url())

    def get_success_url(self):
         return reverse('Homologaciones:home')

    # def get_context_data(self, **kwargs):
    #     context = super().get_context_data(**kwargs)
    #     context['paises'] = queryset=pais.objects.all()
    #     return context
